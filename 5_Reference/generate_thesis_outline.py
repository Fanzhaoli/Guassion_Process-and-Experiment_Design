#!/usr/bin/env python3
"""生成毕业论文大纲 .docx 文件"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

OUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                        "毕业论文大纲_自我优势效应的实验设计空间优化.docx")

doc = Document()

# ── 样式 ──────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = '黑体'
    h.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        h.font.size = Pt(16)
        h.paragraph_format.space_before = Pt(18)
    elif level == 2:
        h.font.size = Pt(14)
        h.paragraph_format.space_before = Pt(12)
    else:
        h.font.size = Pt(12)
        h.paragraph_format.space_before = Pt(8)

# ── 封面 ──────────────────────────────────────────
for _ in range(3):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('硕 士 学 位 论 文 大 纲')
run.font.size = Pt(22)
run.font.name = '黑体'
run.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    '自我优势效应的实验设计空间优化：\n基于Ω的DDM参数预测与验证')
run.font.size = Pt(18)
run.font.name = '黑体'
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Experimental Design Space Optimization of the Self-Prioritization Effect:\n'
          'DDM Parameter Prediction and Validation Based on Ω').font.size = Pt(12)

for _ in range(2):
    doc.add_paragraph()
doc.add_paragraph().add_run('—  2026年  —').font.size = Pt(14)
doc.add_page_break()

# ═══════════════════════════════════════════════════
#  正 文
# ═══════════════════════════════════════════════════

# ── 第1章 ─────────────────────────────────────────
doc.add_heading('第1章  绪论', level=1)

doc.add_heading('1.1  研究背景', level=2)
doc.add_paragraph(
    '自我优势效应（Self-Prioritization Effect, SPE）是指个体对与自我相关的刺激加工更快、'
    '更准的现象（Sui et al., 2012）。在 Self-Matching Task（SMT）范式中，被试通过联结学习'
    '将几何图形与身份标签（自我 vs. 陌生人）建立关联，随后在匹配任务中根据图形-标签的'
    '匹配关系做出按键反应。SMT 范式的一个关键优势在于实验参数（如练习次数 P、刺激呈现'
    '时间 T、反应窗口 W）可以系统性操控，从而直接调控 SPE 的表达。然而，当前文献中缺乏'
    '对不同实验设计参数组合下 SPE 变化的系统性预测框架。本研究旨在构建一个基于漂移扩散'
    '模型（DDM）和高斯过程（GP）的生成模型，以实现对任意实验设计参数组合 (P, T, W) '
    '下 DDM 参数的预测，从而为实验设计优化提供量化工具。')

doc.add_heading('1.2  研究问题与目标', level=2)
doc.add_paragraph('本研究围绕以下核心问题展开：')
items = [
    'RQ1: 遗漏试次（Omission）的不同处理方法如何影响 DDM 参数的估计？（方法论问题）',
    'RQ2: 如何构建一个理论驱动的生成模型，能够从实验设计参数 (P, T, W) 预测 DDM '
    '参数 (v, a, t, z)？（建模问题）',
    'RQ3: 该生成模型的预测能力如何进行交叉验证和行为层面验证？（验证问题）',
    'RQ4: 基于模型的预测不确定性，如何推荐最优的下一轮实验设计参数？（应用问题）',
    'RQ5: 模型对新实验条件的预测能否被独立采集的验证数据所支持？（外部验证问题）',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('1.3  研究意义', level=2)
doc.add_paragraph(
    '理论意义：（1）将 DDM 的参数化预测与高斯过程的非参数灵活性相结合，提供了一种'
    '新的混合建模框架；（2）对 omission 处理方法的系统性比较填补了方法论文献的空白'
    '（回应 Leng et al., 2025 的呼吁）。实践意义：（1）为 SMT 实验设计提供了量化的参数'
    '选择指导；（2）降低了实验设计中的试错成本，提高了 SPE 测量精度。')

doc.add_heading('1.4  论文结构概述', level=2)
doc.add_paragraph('本文共分为三部分：第一部分（第2章）为文献综述，系统回顾 SMT 范式、DDM '
                  '理论及 GP 方法的相关研究；第二部分（第3-5章）包含三项实证研究；第三部分'
                  '（第6章）为总讨论，总结研究发现并提出未来研究方向。')

doc.add_page_break()

# ── 第2章 ─────────────────────────────────────────
doc.add_heading('第2章  文献综述', level=1)

doc.add_heading('2.1  Self-Matching Task 与自我优势效应', level=2)
doc.add_paragraph(
    'Sui et al. (2012) 提出的 SMT 范式通过联结学习建立图形-身份关联，使得实验者可以'
    '精确操控刺激与身份的匹配关系。SPE 的经典指标为自我条件下的正确反应时减去陌生人'
    '条件下的正确反应时（通常为负值，表示自我条件更快）。本节综述 SPE 的行为特征、'
    '影响因素（P, T, W）以及已有的神经机制研究。')

doc.add_heading('2.2  漂移扩散模型（DDM）', level=2)
doc.add_heading('2.2.1  DDM 基本理论与参数', level=3)
doc.add_paragraph(
    '综述 Ratcliff (1978) 及 Ratcliff & McKoon (2008) 的 DDM 理论框架。核心参数包括：'
    '漂移率 v（信息累积速率）、边界分离 a（反应谨慎度）、起始点 z（先验偏向）、非决策'
    '时间 Ter。本节同时引入跨试次变异参数（sv, sz, sTer）及其在模型识别中的作用。')

doc.add_heading('2.2.2  DDM 参数的系统性综述', level=3)
doc.add_paragraph(
    '基于 Tran et al. (2021) 对 158 篇 DDM 文献的系统性定量综述，整理各参数的经验分布、'
    '推荐先验分布及理论边界条件。同时参考 Matzke & Wagenmakers (2009) 的早期参数总结'
    '及三步验证法（仿真-层次贝叶斯-实验操控验证），为本研究的参数校准和模型验证提供'
    '方法论基础。')

doc.add_heading('2.2.3  DDM 中的遗漏试次（Omission）问题', level=3)
doc.add_paragraph(
    '在包含反应截止时间（deadline）的实验设计中，部分试次被试未能做出反应（omission）。'
    '传统做法是直接丢弃这些试次（Drop 方案）或将 RT 设为 deadline 作为截尾数据处理'
    '（Censor 方案）。Leng et al. (2025) 近期通过仿真实验证明，丢弃 omission 会严重偏倚'
    'DDM 参数估计（v 高估、a 低估），即使遗漏率低至 5%。本节综述 omission 建模的'
    '现有方法（截尾法、LAN+OPN 框架、lapse 分布建模）及各自优缺点。')

doc.add_heading('2.3  高斯过程（GP）在认知建模中的应用', level=2)
doc.add_paragraph(
    '综述 Schulz et al. (2018) 的 GP 回归教程及在心理学中的应用。本节重点讨论 GP 在'
    '三个方面的优势：（1）作为代理模型替代显式参数函数；（2）捕捉非线性残差；'
    '（3）提供预测不确定性估计。同时讨论 GP 在小样本条件下的局限性及改进策略。')

doc.add_heading('2.4  实验设计优化', level=2)
doc.add_paragraph(
    '综述 Myung et al. (2013) 的最优实验设计框架及贝叶斯优化方法。讨论基于模型不确定性'
    '的主动学习策略（Active Learning / Bayesian Optimization）如何指导下一轮实验条件的选择。')

doc.add_heading('2.5  文献小结与研究空白', level=2)
doc.add_paragraph(
    '现有文献缺乏一个能从实验设计参数 (P, T, W) 系统性预测 DDM 参数变化的理论框架。'
    '同时，omission 处理方法对 DDM 参数估计的定量影响缺乏来自真实实验数据的验证。'
    '本研究旨在填补这两项空白。')

doc.add_page_break()

# ── 第3章 ─────────────────────────────────────────
doc.add_heading('第3章  研究一：Omission处理方法对DDM参数估计的系统影响', level=1)
doc.add_paragraph('（对应现有 Omission_Sensitivity_Analysis.ipynb 的工作）')

doc.add_heading('3.1  研究问题与假设', level=2)
doc.add_paragraph(
    'RQ1: 在 SMT 实验数据中，Censor（截尾）方案和 Drop（直接丢弃）方案下提取的 '
    'DDM 参数是否存在系统性差异？差异的量级如何随遗漏率变化？')
doc.add_paragraph('假设：遗漏率 >35% 时，两种方案的参数差异将超过可接受范围（95% CI 不重叠），'
                  '验证 Leng et al. (2025) 的仿真结论。')

doc.add_heading('3.2  方法', level=2)
doc.add_heading('3.2.1  被试与实验设计', level=3)
doc.add_paragraph(
    '88 名健康大学生被试（已由课题组前期采集），8 组实验条件，每组 10-12 人。'
    '实验设计空间 Ω = (P, T, W, M)，参数取值范围见表 3-1。')

doc.add_heading('3.2.2  数据处理方案', level=3)
doc.add_paragraph('定义两种 omission 处理方案：')
items2 = [
    'Censor 方案（当前使用）：omission 试次的 RT 设置为 deadline（T+W），response=0，'
    '作为右截尾数据参与 HDDM 层次贝叶斯拟合（p_outlier=0）。',
    'Drop 方案（对照）：直接删除所有 omission 试次，仅使用有效试次进行拟合（p_outlier=0.05）。',
]
for item in items2:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.2.3  HDDM 拟合参数', level=3)
doc.add_paragraph(
    '使用 Docker hcp4715/hddm 镜像。模型设定：depends_on={"v": "identity"}。'
    'MCMC: 3000 draws, 500 burn-in。每组独立拟合。')

doc.add_heading('3.2.4  分析指标', level=3)
doc.add_paragraph(
    '计算两种方案下各 DDM 参数（v_self, v_stranger, SPE_v, a, t, z）的差异 Δ 及其 '
    '95% CI 重叠判断。以 Cohen\'s d 量化效应量。')

doc.add_heading('3.3  结果', level=2)
doc.add_paragraph('（填入 Omission_Sensitivity_Analysis 的具体结果表格和图表）')

doc.add_heading('3.4  讨论', level=2)
doc.add_paragraph(
    '本研究的核心发现包括：（1）G1（遗漏率 72%）的 v_self 在两种方案间出现符号反转'
    '（Censor: v=-4.86 → Drop: v=+1.95），证实了 Leng et al. (2025) 的偏倚机制在真实'
    '数据中的存在；（2）遗漏率 >35% 时，偏倚不可接受（95% CI 不重叠）；'
    '（3）遗漏率 <15% 时，Censor 方案表现足够好。'
    '本研究为后续 DDM 建模时的数据筛选和 omission 处理方法选择提供了实证依据。')

doc.add_page_break()

# ── 第4章 ─────────────────────────────────────────
doc.add_heading('第4章  研究二：SPE实验设计空间的DDM参数化建模与验证', level=1)
doc.add_paragraph('（对应现有 GP+Sigmoid Cleaned Pipeline 的工作）')

doc.add_heading('4.1  研究问题与假设', level=2)
doc.add_paragraph(
    'RQ2: 能否构建一个结合 Sigmoid 理论先验与 GP 残差捕捉的混合生成模型，实现对任意'
    '(P, T, W) 组合下 DDM 参数的可靠预测？')

doc.add_heading('4.2  方法', level=2)
doc.add_heading('4.2.1  实验条件与数据筛选', level=3)
doc.add_paragraph(
    '基于研究一的结论，排除 G1 和 G2（遗漏率 >50%，DDM 拟合不可靠）。G3-G4 标记为'
    '"谨慎使用"，G5-G8 作为核心建模样本（N=6 组）。G7 和 G8 设计重复但被试独立，'
    '均保留为独立数据点。')

doc.add_heading('4.2.2  Sigmoid 理论参数化', level=3)
doc.add_paragraph(
    '定义 Sigmoid v 函数：v(T, P, condition) = v_T(T) × v_P(P) × base_scale_v '
    '× (1 + condition_modulation)。其中 v_T 和 v_P 均为 Sigmoid 形式，condition_modulation '
    '取决于 self（alaph1）还是 stranger（alaph2）条件。定义 Sigmoid a 函数：'
    'a(M) = Sigmoid(M) × base_scale_a × (1 + boundary_modulation)。'
    '共计 7 个可优化参数（alaph1, alaph2, beta1, beta2, gamma, base_scale_v, base_scale_a），'
    '2 个基线参数（T_0=100, k_T=0.01, M_0=600, k_a=0.01 等固定参数）。')

doc.add_heading('4.2.3  Sigmoid 参数校准', level=3)
doc.add_paragraph(
    '使用差分进化算法（Differential Evolution）以最小化 Sigmoid 预测值与 HDDM 真实'
    '参数间的 RMSE 为目标，对 7 个 Sigmoid 参数进行联合优化。搜索边界：alaph1∈[0.01,3.0], '
    'alaph2∈[-2.0,1.0], beta1∈[-2.0,2.0], beta2∈[-1.0,1.0], gamma∈[0.01,1.0], '
    'base_scale_v∈[1.0,10.0], base_scale_a∈[1.0,25.0]。')

doc.add_heading('4.2.4  GP+Sigmoid 混合模型', level=3)
doc.add_paragraph(
    'GPSigmoidHybridModel 架构：对 5 个 DDM 参数（v_self, v_stranger, a, t, z）各训练'
    '一个独立 GP。GP 的核函数为 ConstantKernel × RBF + WhiteKernel。'
    'GP 学习的目标变量为 真实 HDDM 参数 - Sigmoid 预测值 的残差。'
    '最终预测 = Sigmoid 理论预测 + GP 残差预测。对于 t 和 z，Sigmoid 先验取组均值。')

doc.add_heading('4.2.5  验证策略', level=3)
items3 = [
    'In-sample 训练拟合：对比 GP+Sigmoid 预测值与真实 HDDM 参数的 RMSE 和相关性 r。',
    'LOCV（Leave-One-Condition-Out）交叉验证：逐一排除一个条件，在剩余条件上重训练'
    '模型，预测被排除条件的 DDM 参数。',
    '行为层面验证：使用 GP+Sigmoid 预测的 DDM 参数进行 DDM 试次级模拟，对比模拟生成'
    '的行为数据（RT, ACC, Omission Rate, SPE）与真实行为数据的相关性。',
]
for item in items3:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.3  结果', level=2)
doc.add_paragraph('（填入 GP+Sigmoid Cleaned Pipeline 的全部结果表格：'
                  'Sigmoid 校准参数表、训练拟合指标、LOCV 指标、行为验证指标）')

doc.add_heading('4.4  讨论', level=2)
doc.add_paragraph(
    '主要发现：（1）Sigmoid 校准后 alaph1 从默认值 1.5 降至 0.199，说明真实数据中 '
    'self 的相对优势仅为 +20% 而非假设的 +150%；（2）beta1 符号反转（+0.2→-0.826），'
    '提示高 M 条件下决策边界反而降低，与理论假设相反；（3）LOCV 中 v 的 r=-0.09，'
    '说明 6 个训练点不足以支持 GP 在参数层面的泛化；（4）行为层面验证效果出色'
    '（RT r=0.98, ACC r=0.97），证实 Sigmoid 理论先验对行为重建的保障作用。')

doc.add_page_break()

# ── 第5章 ─────────────────────────────────────────
doc.add_heading('第5章  研究三：Omission建模扩展与模型外部验证', level=1)
doc.add_paragraph('【本章为待完成工作，是论文的核心增量贡献部分】')

doc.add_heading('5.1  研究三A：Sigmoid Omission概率函数的建模与校准', level=2)

doc.add_heading('5.1.1  研究动机', level=3)
doc.add_paragraph(
    '研究一已证明 omission 处理方法对参数估计有显著影响，但目前 GP+Sigmoid 混合模型'
    '（研究二）仅预测 DDM 参数 (v, a, t, z)，未对 omission 率本身进行预测。'
    '本研究三A 旨在扩展生成模型，使其能够同时预测实验条件 (P, T, W) 下的 omission 率。')

doc.add_heading('5.1.2  Omission 概率函数的理论参数化', level=3)
doc.add_paragraph(
    '基于实验数据的观察，omission 率受两方面影响：（1）总可用时间 M = T + W——M 越短，'
    '被试越可能来不及反应（遗漏率越高）；（2）练习次数 P——P 越大，被试反应越熟练，'
    '遗漏率越低。因此，omission 概率函数采用双因子 Sigmoid 形式：')
doc.add_paragraph(
    'omission_rate(M, P) = sigmoid_M(M; M_0, k_M) × sigmoid_P(P; P_0, k_P) '
    '× base_scale_omission', style='List Bullet')
doc.add_paragraph(
    '其中 sigmoid_M 捕捉截止时间对遗漏率的主效应：M 越小，遗漏率越高；'
    'sigmoid_P 捕捉练习对遗漏率的调制效应：P 越大，遗漏率越低。')

doc.add_heading('5.1.3  校准方法与评估', level=3)
doc.add_paragraph(
    '使用差分进化算法，以最小化预测遗漏率与观测遗漏率间的 RMSE 为目标，校准 '
    'omission 概率函数的参数。评估指标包括：In-sample RMSE 和 r、LOCV 交叉验证。')

doc.add_heading('5.1.4  与 GP 的整合', level=3)
doc.add_paragraph(
    '在 GPSigmoidHybridModel 中新增 gp_omission（第 6 个 GP），学习 Sigmoid omission '
    '预测值与真实 omission 率之间的残差。最终预测 = Sigmoid omission + GP omission 残差。')

doc.add_heading('5.1.5  DDM 模拟中的 Omission 生成', level=3)
doc.add_paragraph(
    '在行为验证的 DDM 模拟环节（step5），使用预测的 omission 率决定每个试次是否生成'
    'omission：若生成 omission，则该试次的 RT 设为 deadline，response=0。这使得行为'
    '验证不仅对比正确 RT 和 ACC，还能直接对比 omission 率的预测精度。')

doc.add_heading('5.1.6  完成标准', level=3)
doc.add_paragraph('本子研究完成的判据为：')
items4 = [
    'Sigmoid omission 函数的校准 RMSE 低于当前 observed omission 率的组间标准差。',
    'Sigmoid omission + GP omission（如可行）的 LOCV r > 0.5。',
    '整合 omission 后的行为验证中，omission 率预测的 r > 0.85（至少不低于当前 0.92）。',
    '完成第 6 个 GP（gp_omission）的训练、LOCV 和行为验证的更新。',
]
for item in items4:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5.2  研究三B：新实验条件的模型预测验证', level=2)

doc.add_heading('5.2.1  验证实验设计的必要性', level=3)
doc.add_paragraph(
    '研究二的 LOCV 表明，仅靠 6 个训练点的交叉验证不足以评估 GP+Sigmoid 模型的泛化'
    '能力。真正严格的检验是：用现有模型预测未在训练集中出现的新实验条件下的 DDM '
    '参数，然后独立采集该条件下的真实被试数据，对比预测值与实测值。这是模型外部验证'
    '的黄金标准（Myung et al., 2013; Pitt & Myung, 2019）。')

doc.add_heading('5.2.2  验证条件的选择', level=3)
doc.add_paragraph('推荐采集 2-3 个新实验条件。选择原则：')
items5 = [
    '填补设计空间空白：现有 P 仅有 3 个水平（0, 8, 120），需在 9-119 之间选取。',
    '位于 GP 高不确定性区域：T=500ms, W=300-400ms 附近（step6 候选点排名 #1-5）。',
    '实验可行性：确保 T 和 W 取值合理，预期遗漏率 <20%。',
    '构成系统性测试：固定 T/W，变化 P，以检验模型对 P 梯度的预测能力。',
]
for item in items5:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('\n推荐验证条件：')
run.bold = True

# 表格：推荐验证条件
table = doc.add_table(rows=5, cols=7, style='Light Grid Accent 1')
headers = ['条件', 'P', 'T (ms)', 'W (ms)', 'M (ms)', '预期遗漏率', '选择理由']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

rows_data = [
    ['A（必做）', '35', '80', '800', '880', '~15%',
     'P=35 填补 9-119 空白；T/W 与 G7/G8 相同，可检验 P 效应。'],
    ['B（必做）', '60', '80', '800', '880', '~15%',
     'P=60 为额外中间水平；与 A 共同测试 P 梯度。'],
    ['C（推荐）', '35', '500', '350', '850', '~12%',
     '最高 GP 不确定性区域；测试 T/W 泛化。'],
]
for r, row_data in enumerate(rows_data):
    for c, val in enumerate(row_data):
        table.rows[r + 1].cells[c].text = val
        for paragraph in table.rows[r + 1].cells[c].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph()

doc.add_heading('5.2.3  说服导师的论据', level=3)
items6 = [
    '方法论必要性：现有 6 个训练点不足以支持模型泛化声明（LOCV r 普遍为负）。'
    '外部验证是解决这一瓶颈的唯一途径——没有新数据，论文的核心贡献无法被验证。',
    '工作量合理性：仅需 2-3 个新条件（20-36 名被试），使用已有 MATLAB 实验程序，'
    '仅需修改配置文件。预计 2-3 周完成数据采集。',
    '文献支持：Leng et al. (2025), Tran et al. (2021), Myung et al. (2013) 均强调'
    '外部验证在计算建模中的核心地位。缺少验证的建模论文在当前学术标准下难以通过。',
    '与学姐工作的区分：新采集的数据是你独立的贡献，直接回应"你独立完成了什么"的'
    '答辩质疑。',
]
for item in items6:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5.2.4  实验方案', level=3)
doc.add_paragraph(
    '实验程序：使用课题组现有的 MATLAB 实验程序（1_Code/Experiment/exp_matlab/），'
    '修改实验参数 (P, T, W, M)。被试：每组 10-12 名健康大学生（与现有实验相同来源）。'
    '实验时长：每名被试约 20 分钟（260 trials × ~4.5s/trial）。')
doc.add_paragraph(
    'DDM 分析管线：同研究一，使用 Censor 方案（p_outlier=0）。'
    '验证指标：（1）参数层面对比：模型预测的 (v,a,t,z) vs 实测的 (v,a,t,z)；'
    '（2）行为层面对比：模型预测的 (RT, ACC, Omission, SPE) vs 实测值。')

doc.add_heading('5.3  结果（预留）', level=2)
doc.add_paragraph('（填入验证数据的 HDDM 拟合结果及预测 vs 实测对比图表）')

doc.add_heading('5.4  讨论', level=2)
doc.add_paragraph(
    '讨论模型预测的准确性及其边界条件。若预测准确：支持 GP+Sigmoid 作为 SPE 实验设计'
    '优化的有效工具。若预测偏差较大：讨论可能的原因（训练样本过少、Sigmoid 函数形式'
    '限制、实验被试间变异等），提出模型改进方向。')

doc.add_page_break()

# ── 第6章 ─────────────────────────────────────────
doc.add_heading('第6章  总讨论', level=1)

doc.add_heading('6.1  研究发现总结', level=2)
doc.add_paragraph(
    '（1）Omission 处理方法对 DDM 参数估计的系统性影响被定量刻画——遗漏率 >35% 时 '
    '参数偏倚不可接受，遗漏率 <15% 时 Censor 方案足够可靠。'
    '（2）Sigmoid+GP 混合模型在行为重建层面表现优异（RT r=0.98），但参数层面的泛化'
    '受限于训练样本量。'
    '（3）Sigmoid 校准揭示了非平凡的理论参数变化（alaph1≈0.2, beta1 符号反转），'
    '提示真实数据中的 SPE 机制可能与理论假设存在差异。')

doc.add_heading('6.2  理论贡献', level=2)
items7 = [
    '提出了 Sigmoid（理论先验）+ GP（数据驱动残差）的混合建模框架，在认知建模领域'
    '具有一定的方法论创新性。',
    '通过 Omission 敏感性分析，为 DDM 应用中的 omission 处理方法选择提供了实证指导。',
    '系统收集并整合了 DDM 参数的经验分布数据（Tran et al., 2021; Matzke & Wagenmakers, '
    '2009），建立了面向本项目 SMT 范式的参数参考标准。',
]
for item in items7:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.3  方法论贡献', level=2)
items8 = [
    '开发了完整的 DDM 分析工具链（Docker HDDM 拟合 → 参数提取 → Sigmoid 校准 → '
    'GP 建模 → LOCV → 行为验证 → 候选推荐），可作为类似研究的方法论模板。',
    '将 Omission 概率函数集成到生成模型中，首次实现了对 SPE 实验中遗漏率的定量预测。',
]
for item in items8:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.4  局限与未来方向', level=2)
items9 = [
    '样本量限制：仅 6-8 个设计条件限制了 GP 的泛化能力。未来研究应系统性增加 '
    '实验条件数量（尤其 P 的中间水平）。',
    'Sigmoid 函数形式：base_scale_a 触及搜索上界，提示 a 的 Sigmoid 参数化'
    '可能需调整为更灵活的函数形式（如线性组合或二次模型）。',
    'DDM 模型选择：当前使用恒定边界 DDM。未来可考虑塌缩边界模型（ANGLE/WEIBULL）'
    '以更好地捕捉 deadline 下的决策动态（Leng et al., 2025）。',
    '被试间变异：G7/G8 设计相同但参数不同，提示被试间变异可能大于实验操控效应。'
    '可考虑在 GP 建模中纳入 hierarchical 结构。',
    'NonMatching 试次：当前仅分析 Matching 试次。NonMatching 条件下的决策过程'
    '可能提供额外的模型约束信息。',
]
for item in items9:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ── 参考文献 ──────────────────────────────────────
doc.add_heading('参考文献', level=1)
refs = [
    'Leng, X., Fengler, A., Shenhav, A., & Frank, M. J. (2025). The Perils of Omitting '
    'Omissions when Modeling Evidence Accumulation. In Prep.',
    'Matzke, D., & Wagenmakers, E.-J. (2009). Psychological interpretation of the '
    'ex-Gaussian and shifted Wald parameters: A diffusion model analysis. '
    'Psychonomic Bulletin & Review, 16(5), 798-817.',
    'Myung, J. I., Cavagnaro, D. R., & Pitt, M. A. (2013). A tutorial on adaptive '
    'design optimization. Journal of Mathematical Psychology, 57(3), 53-67.',
    'Pitt, M. A., & Myung, J. I. (2019). Robust modeling through design optimization. '
    'Computational Brain & Behavior, 2, 200-201.',
    'Ratcliff, R. (1978). A theory of memory retrieval. Psychological Review, 85, 59-108.',
    'Ratcliff, R., & McKoon, G. (2008). The diffusion decision model: Theory and data '
    'for two-choice decision tasks. Neural Computation, 20, 873-922.',
    'Schulz, E., Speekenbrink, M., & Krause, A. (2018). A tutorial on Gaussian process '
    'regression: Modelling, exploring, and exploiting functions. Journal of Mathematical '
    'Psychology, 85, 1-16.',
    'Sui, J., He, X., & Humphreys, G. W. (2012). Perceptual effects of social salience: '
    'Evidence from self-prioritization effects on perceptual matching. Journal of '
    'Experimental Psychology: Human Perception and Performance, 38(5), 1105-1117.',
    'Tran, N. H., van Maanen, L., Heathcote, A., & Matzke, D. (2021). Systematic '
    'Parameter Reviews in Cognitive Modeling: Towards a Robust and Cumulative '
    'Characterization of Psychological Processes in the Diffusion Decision Model. '
    'Frontiers in Psychology, 11, 608287.',
    'van Ravenzwaaij, D., Dutilh, G., & Wagenmakers, E.-J. (2011). Cognitive model '
    'decomposition of the BART: Assessment and application. Journal of Mathematical '
    'Psychology, 55, 94-105.',
    'Wiecki, T. V., Sofer, I., & Frank, M. J. (2013). HDDM: Hierarchical Bayesian '
    'estimation of the Drift-Diffusion Model in Python. Frontiers in Neuroinformatics, '
    '7, 14.',
]
for i, ref in enumerate(refs):
    doc.add_paragraph(f'[{i+1}] {ref}', style='List Number')

doc.add_page_break()

# ── 附录 ──────────────────────────────────────────
doc.add_heading('附录', level=1)

doc.add_heading('附录A：Sigmoid 参数全表', level=2)
doc.add_paragraph('（引用 5_Reference/RoadMap.md 的 Phase 0.4 内容，列出所有显式及隐性参数）')

doc.add_heading('附录B：DDM 参数参考手册', level=2)
doc.add_paragraph('（引用 5_Reference/RoadMap.md 的附录章节，含参数范围规范和异常值剔除标准）')

doc.add_heading('附录C：Omission 建模可行性分析', level=2)
doc.add_paragraph('（引用 5_Reference/Omission建模可行性分析.md）')

doc.add_heading('附录D：推荐验证实验的详细方案', level=2)
doc.add_paragraph('包含：被试招募方案、实验程序修改说明、预期数据量估算、时间节点规划。')

doc.add_heading('附录E：代码与数据可复现性说明', level=2)
doc.add_paragraph(
    '项目完整代码位于 1_Code/，数据位于 2_Data/，图表位于 3_Figures/。'
    '所有分析均可通过运行对应 Python 脚本复现。计划在论文发表前将代码和数据上传至 '
    'OSF 或 GitHub 开源仓库。')

# ── 保存 ──────────────────────────────────────────
doc.save(OUT_PATH)
print(f'✅ 毕业论文大纲已保存到: {OUT_PATH}')
